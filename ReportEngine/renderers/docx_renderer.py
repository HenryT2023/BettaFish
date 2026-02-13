# -*- coding: utf-8 -*-
"""
DOCX Renderer — 将 ReportEngine IR 渲染为 .docx 文件

输出兼容 wechat_publisher/docx_parser.py 的格式。
支持的 IR block 类型：heading, paragraph, list, blockquote, hr, table, callout。
不支持（WeChat 不兼容）：math, code, figure, chart, widget。
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    raise ImportError("python-docx 未安装，请运行: pip install python-docx")


class DocxRenderer:
    """将 IR blocks 渲染为 .docx 文件"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """配置基本样式，确保中文字体"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(11)

        # 设置中文字体（通过 XML 操作）
        try:
            from docx.oxml.ns import qn
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        except Exception:
            pass

    def render_blocks(self, blocks: List[Dict[str, Any]]) -> Document:
        """渲染 IR blocks 列表"""
        for block in blocks:
            block_type = block.get("type", "")
            handler = getattr(self, f"_render_{block_type}", None)
            if handler:
                handler(block)
            else:
                logger.debug(f"跳过不支持的 block 类型: {block_type}")
        return self.doc

    def render_from_markdown(self, markdown_text: str) -> Document:
        """
        从 Markdown 文本渲染 .docx（简化路径，不走 IR）。
        适用于 Quill 直接从 LLM 输出 Markdown 的场景。
        """
        lines = markdown_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # 空行
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith("#"):
                level = 0
                while level < len(line) and line[level] == "#":
                    level += 1
                text = line[level:].strip()
                self._add_heading(text, min(level, 4))
                i += 1
                continue

            # 分隔线
            if line.strip() in ("---", "***", "___"):
                self._render_hr({})
                i += 1
                continue

            # 引用
            if line.startswith(">"):
                quote_lines = []
                while i < len(lines) and lines[i].startswith(">"):
                    quote_lines.append(lines[i].lstrip("> ").rstrip())
                    i += 1
                self._add_blockquote("\n".join(quote_lines))
                continue

            # 无序列表
            if line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
                list_items = []
                while i < len(lines):
                    cur = lines[i].lstrip()
                    if cur.startswith("- ") or cur.startswith("* "):
                        item_text = cur[2:].strip() if len(cur) > 2 else ""
                        if item_text:
                            list_items.append(item_text)
                        i += 1
                    elif not lines[i].strip():
                        # 空行：如果下一个非空行仍是列表项，则跳过继续
                        peek = i + 1
                        while peek < len(lines) and not lines[peek].strip():
                            peek += 1
                        if peek < len(lines) and (lines[peek].lstrip().startswith("- ") or lines[peek].lstrip().startswith("* ")):
                            i += 1
                        else:
                            break
                    else:
                        break
                for item in list_items:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_rich_text(p, item)
                continue

            # 有序列表
            if len(line.lstrip()) > 2 and line.lstrip()[0].isdigit() and ". " in line.lstrip()[:5]:
                list_items = []
                while i < len(lines):
                    stripped = lines[i].lstrip()
                    if len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:5]:
                        item_text = stripped.split(". ", 1)[-1].strip()
                        if item_text:
                            list_items.append(item_text)
                        i += 1
                    elif not lines[i].strip():
                        # 空行：如果下一个非空行仍是有序列表项，则跳过继续
                        peek = i + 1
                        while peek < len(lines) and not lines[peek].strip():
                            peek += 1
                        if peek < len(lines) and len(lines[peek].lstrip()) > 2 and lines[peek].lstrip()[0].isdigit() and ". " in lines[peek].lstrip()[:5]:
                            i += 1
                        else:
                            break
                    else:
                        break
                for item in list_items:
                    p = self.doc.add_paragraph(style="List Number")
                    self._add_rich_text(p, item)
                continue

            # Markdown 表格（| col1 | col2 | 格式）
            if line.strip().startswith("|") and "|" in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._render_markdown_table(table_lines)
                continue

            # 普通段落
            p = self.doc.add_paragraph()
            self._add_rich_text(p, line)
            i += 1

        return self.doc

    def insert_images(self, image_paths: List[str], position: str = "end"):
        """
        按图片类型智能插入到文章对应位置：
        - cover → 第一个 H1 标题之后
        - trend → 第一个 H2 之后（正文起始处）
        - gap / gap-ai → 含"信息差"/"数据"/"对比"的 H2 之后
        - 其他 → 文末兜底
        """
        if not image_paths:
            return

        # 按文件名分类
        cover, trend, gap_imgs, others = [], [], [], []
        for p in image_paths:
            if not Path(p).exists():
                logger.warning(f"图片不存在，跳过: {p}")
                continue
            name = Path(p).stem.lower()
            if "cover" in name:
                cover.append(p)
            elif "trend" in name:
                trend.append(p)
            elif "gap" in name:
                gap_imgs.append(p)
            else:
                others.append(p)

        paragraphs = self.doc.paragraphs

        # 找锚点段落索引
        first_h1_idx = None
        first_h2_idx = None
        gap_h2_idx = None
        last_h2_idx = None
        gap_keywords = ["信息差", "数据", "对比", "海外", "国内", "gap"]

        for i, para in enumerate(paragraphs):
            style_name = (para.style.name or "").lower()
            is_heading = "heading" in style_name
            if not is_heading:
                continue
            text_lower = para.text.lower()
            if "heading 1" in style_name and first_h1_idx is None:
                first_h1_idx = i
            if "heading 2" in style_name:
                if first_h2_idx is None:
                    first_h2_idx = i
                last_h2_idx = i
                if gap_h2_idx is None and any(kw in text_lower for kw in gap_keywords):
                    gap_h2_idx = i

        # gap 图兜底：如果没找到匹配的 H2，放在倒数第二个 H2 后
        if gap_h2_idx is None and last_h2_idx is not None:
            gap_h2_idx = last_h2_idx

        # 插入（从后往前插，避免索引偏移）
        insertions = []  # [(paragraph_index, [image_paths])]

        if cover and first_h1_idx is not None:
            insertions.append((first_h1_idx, cover))
        elif cover:
            insertions.append((0, cover))

        if trend and first_h2_idx is not None:
            insertions.append((first_h2_idx, trend))

        if gap_imgs and gap_h2_idx is not None:
            insertions.append((gap_h2_idx, gap_imgs))

        # 从后往前排序，后面的先插入
        insertions.sort(key=lambda x: -x[0])

        for para_idx, imgs in insertions:
            self._insert_images_after_paragraph(paragraphs[para_idx], imgs)

        # 兜底：未匹配的图片 append 到文末
        for img_path in others:
            self._append_image_at_end(img_path)

    def _insert_images_after_paragraph(self, anchor_para, image_paths: List[str]):
        """在指定段落之后插入图片（操作底层 XML）"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
        from docx.shared import Emu

        # 从后往前插入以保持顺序
        for img_path in reversed(image_paths):
            try:
                # 创建新段落元素
                new_para = OxmlElement("w:p")
                # 居中对齐
                pPr = OxmlElement("w:pPr")
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                pPr.append(jc)
                new_para.append(pPr)

                # 创建 run
                run_elem = OxmlElement("w:r")
                new_para.append(run_elem)

                # 在锚点段落之后插入空段落 + 图片段落
                anchor_para._element.addnext(new_para)

                # 用 python-docx 的 add_picture 方式添加图片到 run
                from docx.shared import Inches as _Inches
                # 通过临时段落方式获取图片 relationship
                tmp_para = self.doc.add_paragraph()
                run = tmp_para.add_run()
                run.add_picture(img_path, width=_Inches(5.5))
                tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 把临时段落的 XML 移到正确位置
                anchor_para._element.addnext(tmp_para._element)

                # 删除之前插入的空 new_para
                new_para.getparent().remove(new_para)

                logger.info(f"图片已插入 DOCX（就近）: {Path(img_path).name}")
            except Exception as e:
                logger.warning(f"就近插入图片失败 {img_path}: {e}")
                # 降级到文末
                self._append_image_at_end(img_path)

    def _append_image_at_end(self, img_path: str):
        """兜底：在文末追加图片"""
        try:
            self.doc.add_paragraph()
            self.doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logger.info(f"图片已插入 DOCX（文末）: {Path(img_path).name}")
        except Exception as e:
            logger.warning(f"插入图片失败 {img_path}: {e}")

    def save(self, output_path: str) -> str:
        """保存 .docx 文件"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"DOCX 已保存: {output_path}")
        return output_path

    # ========== IR Block Handlers ==========

    def _render_heading(self, block: Dict):
        level = block.get("level", 1)
        text = block.get("text", "")
        self._add_heading(text, level)

    def _render_paragraph(self, block: Dict):
        p = self.doc.add_paragraph()
        inlines = block.get("inlines", [])
        if inlines:
            for inline in inlines:
                self._add_inline_run(p, inline)
        else:
            p.add_run("")

        align = block.get("align")
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _render_list(self, block: Dict):
        list_type = block.get("listType", "bullet")
        items = block.get("items", [])
        style_name = "List Bullet" if list_type == "bullet" else "List Number"

        for item_blocks in items:
            p = self.doc.add_paragraph(style=style_name)
            for sub_block in item_blocks:
                if sub_block.get("type") == "paragraph":
                    for inline in sub_block.get("inlines", []):
                        self._add_inline_run(p, inline)

    def _render_blockquote(self, block: Dict):
        inner_blocks = block.get("blocks", [])
        text_parts = []
        for b in inner_blocks:
            if b.get("type") == "paragraph":
                for inline in b.get("inlines", []):
                    text_parts.append(inline.get("text", ""))
        self._add_blockquote(" ".join(text_parts))

    def _render_engineQuote(self, block: Dict):
        title = block.get("title", "Agent 观点")
        inner_blocks = block.get("blocks", [])
        text_parts = []
        for b in inner_blocks:
            if b.get("type") == "paragraph":
                for inline in b.get("inlines", []):
                    text_parts.append(inline.get("text", ""))
        self._add_blockquote(f"[{title}] " + " ".join(text_parts))

    def _render_hr(self, block: Dict):
        p = self.doc.add_paragraph()
        p.add_run("─" * 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_table(self, block: Dict):
        rows_data = block.get("rows", [])
        if not rows_data:
            return

        max_cols = max(len(r.get("cells", [])) for r in rows_data)
        table = self.doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = "Table Grid"

        for row_idx, row in enumerate(rows_data):
            for col_idx, cell in enumerate(row.get("cells", [])):
                if col_idx < max_cols:
                    cell_blocks = cell.get("blocks", [])
                    text = ""
                    for b in cell_blocks:
                        if b.get("type") == "paragraph":
                            for inline in b.get("inlines", []):
                                text += inline.get("text", "")
                    table.cell(row_idx, col_idx).text = text

    def _render_callout(self, block: Dict):
        tone = block.get("tone", "info")
        title = block.get("title", "")
        prefix_map = {"info": "💡", "warning": "⚠️", "success": "✅", "danger": "❌"}
        prefix = prefix_map.get(tone, "📌")

        inner_blocks = block.get("blocks", [])
        text_parts = []
        for b in inner_blocks:
            if b.get("type") == "paragraph":
                for inline in b.get("inlines", []):
                    text_parts.append(inline.get("text", ""))

        full_text = f"{prefix} {title}\n{' '.join(text_parts)}" if title else f"{prefix} {' '.join(text_parts)}"
        self._add_blockquote(full_text)

    def _render_kpiGrid(self, block: Dict):
        items = block.get("items", [])
        if not items:
            return
        table = self.doc.add_table(rows=2, cols=len(items))
        table.style = "Table Grid"
        for i, item in enumerate(items):
            table.cell(0, i).text = item.get("label", "")
            value_text = item.get("value", "")
            if item.get("unit"):
                value_text += f" {item['unit']}"
            table.cell(1, i).text = value_text

    # ========== Helpers ==========

    def _render_markdown_table(self, table_lines: List[str]):
        """解析 Markdown 表格（| col | col |）并渲染为 docx 表格"""
        if len(table_lines) < 2:
            return

        def parse_row(line: str) -> List[str]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            return cells

        # 跳过分隔行（|---|---|）
        data_lines = []
        for line in table_lines:
            stripped = line.replace("|", "").replace("-", "").replace(":", "").strip()
            if stripped:  # 非纯分隔行
                data_lines.append(line)

        if not data_lines:
            return

        rows = [parse_row(line) for line in data_lines]
        max_cols = max(len(r) for r in rows)

        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for row_idx, row in enumerate(rows):
            for col_idx, cell_text in enumerate(row):
                if col_idx < max_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    self._add_rich_text(p, cell_text)
                    # 表头加粗 + 底色
                    if row_idx == 0:
                        for run in p.runs:
                            run.bold = True
                        from docx.oxml.ns import qn
                        from docx.oxml import OxmlElement
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "2C3E50")
                        shading.set(qn("w:val"), "clear")
                        cell._tc.get_or_add_tcPr().append(shading)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)

    def _add_heading(self, text: str, level: int):
        heading = self.doc.add_heading(text, level=min(level, 4))
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_blockquote(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    def _add_inline_run(self, paragraph, inline: Dict):
        """添加 IR inline run 到段落"""
        text = inline.get("text", "")
        run = paragraph.add_run(text)

        marks = inline.get("marks", [])
        for mark in marks:
            mark_type = mark.get("type", "")
            if mark_type == "bold":
                run.bold = True
            elif mark_type == "italic":
                run.italic = True
            elif mark_type == "underline":
                run.underline = True
            elif mark_type == "strike":
                run.font.strike = True
            elif mark_type == "color":
                color_val = mark.get("value", "")
                if isinstance(color_val, str) and len(color_val) == 7 and color_val.startswith("#"):
                    try:
                        run.font.color.rgb = RGBColor(
                            int(color_val[1:3], 16),
                            int(color_val[3:5], 16),
                            int(color_val[5:7], 16),
                        )
                    except (ValueError, IndexError):
                        pass

    def _add_rich_text(self, paragraph, text: str):
        """解析简单 Markdown 标记（**bold**, *italic*）并添加到段落"""
        import re
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
